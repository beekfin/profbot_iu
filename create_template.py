from docx import Document

def create_template():
    doc = Document()
    doc.add_heading('Заявление на материальную помощь', 0)

    doc.add_paragraph('Председателю профбюро факультета ИУ')
    doc.add_paragraph('От студента группы {{group}}')
    doc.add_paragraph('{{fio}}')
    
    doc.add_paragraph('')
    doc.add_paragraph('Заявление', style='Title')
    
    doc.add_paragraph(
        'Прошу оказать мне материальную помощь в связи с тяжелым материальным положением.'
    )
    
    doc.add_paragraph('')
    doc.add_paragraph('Дата: {{date}}')
    doc.add_paragraph('Подпись: _________________')

    doc.save('app/templates/material_aid.docx')
    print("Template created.")

if __name__ == "__main__":
    create_template()
